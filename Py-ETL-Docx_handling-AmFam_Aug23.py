#%%

#Author: Chinmay Patane, Product Owner/Business Analyst Lead, American Family Mutual Insurance Co., Built for Commercial Auto Line of Business in Pricing

##Details: This ETL is written in Python. At AmFam commercial Insurance Company, each month new enterprise rates need to be uploaded for every LoB 
#(Commericial Auto, Worker's Compensation, Business Owners, Commercial Property, General Liability etc.) so that competitive market rate 
#insurance premium can be offered to the customers. 

##This ETL is made for LoB commercial Auto (CA) when certain systems are still using the legacy .xls files,
#(like BaseRate,Primaryt, Secondary etc). This ETL is made with a purpose of scanning the requirement contents in docx file. In some cases, online pages are
#not accesible because of API failure/not available. In such cases, this ETL reads data from exported web page in docx form to proceed with: 

#Read the requirements from .docx format and extract the data from it, read the expected data an convert to tabular form
#Create a report to notify stakeholders - consisting of how many states with their effective dates are modofied. This report is used in Unit Testing and User Acceptance Testing 
#Create a report to make sure all the given requirements match data extracted and ready to be used in rating

#%%
#Initiate the ETL. Please enter your desired release number and product version
Revision_Number = '23.3.20'
Product_Version = '23060600'
Path_to_save_files = 'C:/Users/chinp/OneDrive/Desktop/Python/Currently Building/'
Requirements_Path = 'C:/Users/chinp/OneDrive/Desktop/Python/Currently Building/August 2023.docx'

#%%
#Import libraries
import pandas as pd
import openpyxl
from docx import Document
from openpyxl import load_workbook
from datetime import datetime
#%%
#Separate the tables and paragraphs from the table, and then get the required table using that table's index
def read_word_document(Requirements_Path):
    doc = Document(Requirements_Path)
    paragraphs = []
    tables = []
    
    # Iterate over paragraphs in the document
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)
    
    # Iterate over tables in the document
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            cell_data = []
            for cell in row.cells:
                cell_data.append(cell.text)
            table_data.append(cell_data)
        tables.append(table_data)
    
    return paragraphs, tables

# Call the function to read the Word document
paragraphs, tables = read_word_document(Requirements_Path) # Read whole document

Get_Requirements = pd.DataFrame(tables[9])
Get_Requirements.columns = Get_Requirements.iloc[0]
Get_Requirements = Get_Requirements[1:]

Get_Requirements.columns = [
    "State",
    "NB_Effective_Date",
    "RN_Effective_Date",
    "Rate_Table_Changes",
    "REVISION"
]

# Convert date columns to datetime
Get_Requirements['NB_Effective_Date'] = pd.to_datetime(Get_Requirements['NB_Effective_Date'], errors='coerce')
Get_Requirements['RN_Effective_Date'] = pd.to_datetime(Get_Requirements['RN_Effective_Date'], errors='coerce')


Get_Requirements['NB_Effective_Date'] = Get_Requirements['NB_Effective_Date'].dt.strftime('%m/%d/%Y')
Get_Requirements['RN_Effective_Date'] = Get_Requirements['RN_Effective_Date'].dt.strftime('%m/%d/%Y')

#%%
# Initiate the report
Requirements = pd.DataFrame(columns=[
    "Rate_Table_Changes",
    "IEL_Table",
    "States",
    "States_Total",
    "Revisions",
    "Number_of_Lines",
    "Revisions_Found",
    "States_Found",
    "States_Found_Total",
    "Mismatch_Found"
])

# Extract unique Rate_Table_Changes values
unique_rates = Get_Requirements['Rate_Table_Changes'].unique()


for rate in unique_rates:
    states = Get_Requirements.loc[Get_Requirements['Rate_Table_Changes'] == rate, 'State']
    formatted_dates = pd.to_datetime(Get_Requirements.loc[Get_Requirements['Rate_Table_Changes'] == rate, 'NB_Effective_Date'], errors='coerce').dt.strftime('%m/%d/%Y')
    states_with_dates = [f"{state} ({date})" for state, date in zip(states, formatted_dates)]
    states_string = ', '.join(states_with_dates)
    
    revisions = Get_Requirements.loc[Get_Requirements['Rate_Table_Changes'] == rate, 'REVISION'].unique()
    revisions_string = ', '.join(revisions)
    
    new_row = pd.DataFrame({
        'Rate_Table_Changes': [rate],
        'States': [states_string],
        'Revisions': [revisions_string]
    })
    
    Requirements = pd.concat([Requirements, new_row], ignore_index=True)



# Export the Requirements DataFrame to Excel
Requirements['States_Total'] = Requirements['States'].apply(lambda x: len(set(map(str.strip, x.split(',')))))

#%%

#Read the tabs from Rating Matrix which are in the requirement
required_tabs = Requirements["Rate_Table_Changes"].tolist()
Data1 = pd.read_excel("C:/Users/chinp/OneDrive/Desktop/Python/Currently Building/CA-Rating-Algorithm-Matrix.xlsx", sheet_name=required_tabs)

# Get mapping of factors vs tables
Mapping = pd.read_excel("C:/Users/chinp/OneDrive/Desktop/Python/Currently Building/Mapping.xlsx")

# Filter mapping based on required factors
Mapping = Mapping[Mapping["Factor"].isin(required_tabs)]

# Re-order mapping to match the datasets
Mapping = Mapping.set_index("Factor").loc[required_tabs].reset_index()

#%%
#Export the data and create the report in one for loop

Factor_Name = Mapping["Factor"].tolist()
Table_Name = Mapping["Table"].tolist()

# Export Table and Report
for i in range(len(Factor_Name)):
    f = Factor_Name[i]
    t = Table_Name[i]
    
    Data3 = Data1[f].copy()
    Data3 = pd.DataFrame(Data3)
    Data3.columns = Data3.iloc[0]
    Data3 = Data3[1:]
    
    Data3 = Data3.dropna(subset=["REVISION"])  # Get only those with revision number mentioned
    Data3 = Data3[Data3["REVISION"] == Revision_Number]  # Get desired revision number for this release
    
    Data3["PRODUCT_VERSION"] = Product_Version  # Desired Product Version
    Data3["SOURCE"] = "Custom"  # Source is always Custom
    Data3 = Data3.loc[:, ~Data3.columns.isin(["RESOURCE_OWNER", "CUSTOMER_CODE", "RELEASE_NUMBER", "GID"])]
    #Export the data to upload thereafter
    Data3.to_excel(f"{Path_to_save_files}{t}.xlsx", index=False)
    
    # Make Report for whether given requirements match with the matrix data
    Requirements.loc[i, "Number_of_Lines"] = len(Data3)
    Requirements.loc[i, "Revisions_Found"] = ", ".join(Data3["REVISION"].unique())
    
    unique_states_dates = [f"{state} ({date})" for state, date in zip(Data3["STATE_CODE"], Data3["EFFECTIVE_DATE"])]
    unique_states_dates = pd.Index(unique_states_dates).unique()
    Requirements.loc[i, "States_Found"] = ", ".join(unique_states_dates)
    
    Requirements.loc[i, 'IEL_Table'] = t

    
    # Fill the last column, Mismatch_Found, in the report based on the states mentioned in the requirements vs states found in the rating matrix mentioned in States_Found
    X = [x.strip() for x in str(Requirements.loc[i, "States"]).split(",")]
    Y = [y.strip() for y in str(Requirements.loc[i, "States_Found"]).split(",")]
    
    if set(X) == set(Y):
        Requirements.loc[i, "Mismatch_Found"] = "No"
    else:
        Requirements.loc[i, "Mismatch_Found"] = "Yes"

Requirements["States_Found_Total"] = Requirements["States_Found"].str.split(",").apply(lambda x: len(set([state.strip() for state in x])))

#Export the report
Requirements.to_excel(f"{Path_to_save_files}Requirements_Report.xlsx", index=False)
#%%














